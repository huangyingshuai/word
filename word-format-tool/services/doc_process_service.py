"""
模块名称：services.doc_process_service
模块职责：文档排版全流程业务封装
"""
import copy
import docx
from docx.shared import Cm, Pt
from typing import Tuple, Dict, Optional, List
from config.constants import ALIGN_MAP, LINE_TYPE_MAP, FONT_SIZE_NUM
from config.templates import validate_template
from utils.doc_utils import is_protected_para, set_run_font, process_number_in_para
from utils.file_utils import safe_save_temp_file, safe_delete_file, clear_memory
from core.title_recognizer import get_title_level_with_context
from core.number_recognizer import identify_number_item
from core.number_grouper import group_number_items, convert_number_format

# ====================== 核心业务服务 ======================
def process_document(
    file_content: bytes,
    config: Dict,
    number_config: Dict,
    enable_title_regex: bool = True,
    enable_context_check: bool = True,
    force_style: bool = True,
    keep_spacing: bool = False,
    clear_blank: bool = False,
    max_blank: int = 1,
    corrected_number_items: Optional[List] = None,
    group_format_config: Optional[Dict] = None,
    disable_auto_numbering: bool = True,  # 新增：默认禁用自动编号
    target_number_format: str = "1."       # 新增：序号统一格式
) -> Tuple[Optional[bytes], Optional[Dict], Optional[List]]:
    tmp_path = None
    output_path = None
    doc = None
    try:
        is_valid, msg = validate_template(config)
        if not is_valid:
            raise ValueError(f"格式配置校验失败：{msg}")

        tmp_path = safe_save_temp_file(file_content)
        if not tmp_path:
            raise Exception("临时文件保存失败")

        doc = docx.Document(tmp_path)
        stats = {"一级标题":0,"二级标题":0,"三级标题":0,"正文":0,"表格":0,"图片":0,"序号项":0}
        title_records = []
        total_para = len(doc.paragraphs)

        # 图片统计
        original_image_count = 0
        for para in doc.paragraphs:
            try:
                original_image_count += len(para._element.findall('.//w:drawing', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}))
                original_image_count += len(para._element.findall('.//w:pict', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}))
            except:
                pass
        stats["图片"] = original_image_count

        # 【新增】彻底禁用Word自动编号
        if disable_auto_numbering:
            # 1. 清理所有段落的编号格式
            for para in doc.paragraphs:
                try:
                    para.paragraph_format.numbering_style = None
                    # 直接操作XML移除编号属性
                    if para._element.pPr is not None:
                        numPr = para._element.pPr.find('.//w:numPr', 
                            namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                        if numPr is not None:
                            para._element.pPr.remove(numPr)
                except Exception:
                    pass
            
            # 2. 清理所有样式的编号关联
            for style in doc.styles:
                try:
                    if hasattr(style, 'paragraph_format'):
                        style.paragraph_format.numbering_style = None
                except Exception:
                    pass

        # 预扫描
        pre_scan_result = [ {"type": "protected"} for _ in range(total_para) ]
        number_items = []
        title_result = get_title_level_with_context(doc.paragraphs, enable_regex=enable_title_regex, enable_context_check=enable_context_check)

        for para_idx, para in enumerate(doc.paragraphs):
            if is_protected_para(para):
                continue
            _, title_level = title_result[para_idx]
            number_info = identify_number_item(para, para_idx)
            if number_info:
                number_items.append(number_info)
            pre_scan_result[para_idx] = {
                "type": "normal",
                "title_level": title_level,
                "number_info": number_info
            }
            # 修复：移除海象运算符
            text_content = para.text.strip()
            if title_level in stats:
                stats[title_level] += 1
                title_records.append({
                    "段落序号": para_idx,
                    "识别结果": title_level,
                    "文本内容": text_content[:50] + "..." if len(text_content) > 50 else text_content
                })
            else:
                stats["正文"] += 1

        # 序号分组与转换
        if corrected_number_items is not None:
            number_items = copy.deepcopy(corrected_number_items)
        number_groups = group_number_items(number_items)
        stats["序号项"] = len(number_items)

        # 【新增】序号统一转换
        converted_all_items = []
        for group in number_groups:
            converted_items = convert_number_format(group, target_number_format)
            converted_all_items.extend(converted_items)
        
        # 【新增】将转换后的序号应用到文档
        for item in converted_all_items:
            para_idx = item["para_index"]
            if 0 <= para_idx < len(doc.paragraphs):
                para = doc.paragraphs[para_idx]
                if not is_protected_para(para):
                    para.text = item["full_text"]

        # 样式绑定
        if force_style:
            required_styles = [("标题 1", "Heading 1", 1), ("标题 2", "Heading 2", 2), ("标题 3", "Heading 3", 3), ("正文", "Normal", 0)]
            for cn_name, en_name, outline_level in required_styles:
                try:
                    if cn_name not in doc.styles:
                        doc.styles.add_style(cn_name, 1)
                        doc.styles[cn_name].base_style = doc.styles[en_name]
                except Exception:
                    pass
            for level in range(1,10):
                try:
                    doc.styles.add_style(f"列表{level}级", 1)
                except Exception:
                    pass

        # 格式应用
        for para_idx, para in enumerate(doc.paragraphs):
            scan_info = pre_scan_result[para_idx]
            if scan_info["type"] == "protected":
                continue
            title_level = scan_info["title_level"]
            number_info = scan_info["number_info"]

            # 样式
            if force_style:
                try:
                    if title_level == "一级标题": para.style = "标题 1"
                    elif title_level == "二级标题": para.style = "标题 2"
                    elif title_level == "三级标题": para.style = "标题 3"
                    elif number_info: para.style = f"列表{number_info['level']}级"
                    else: para.style = "正文"
                except Exception:
                    pass

            # 配置获取
            cfg = config[title_level] if not number_info else config["序号列表"]
            font_size = FONT_SIZE_NUM[cfg["size"]]

            # 段落格式（修复：单位Cm→Pt）
            try:
                if ALIGN_MAP[cfg["align"]]:
                    para.alignment = ALIGN_MAP[cfg["align"]]
                para.paragraph_format.line_spacing_rule = LINE_TYPE_MAP[cfg["line_type"]]
                if cfg["line_type"] == "多倍行距":
                    para.paragraph_format.line_spacing = cfg["line_value"]
                elif cfg["line_type"] == "固定值":
                    para.paragraph_format.line_spacing = Pt(cfg["line_value"])
                if not keep_spacing:
                    para.paragraph_format.space_before = Pt(cfg.get("space_before", 0))
                    para.paragraph_format.space_after = Pt(cfg.get("space_after", 0))
                if title_level == "正文" and cfg["indent"] > 0:
                    para.paragraph_format.first_line_indent = Cm(cfg["indent"] * 0.37)
            except Exception:
                continue

            # 字体
            try:
                if title_level == "正文" and not number_info:
                    process_number_in_para(para, cfg["font"], font_size, number_config)
                else:
                    for run in para.runs:
                        set_run_font(run, cfg["font"], font_size, cfg["bold"])
            except Exception:
                continue

        # 表格处理
        for table in doc.tables:
            stats["表格"] +=1
            cfg = config["表格"]
            font_size = FONT_SIZE_NUM[cfg["size"]]
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        try:
                            if ALIGN_MAP[cfg["align"]]: p.alignment = ALIGN_MAP[cfg["align"]]
                            p.paragraph_format.line_spacing_rule = LINE_TYPE_MAP[cfg["line_type"]]
                            if cfg["line_type"] == "固定值": p.paragraph_format.line_spacing = Pt(cfg["line_value"])
                            for run in p.runs:
                                set_run_font(run, cfg["font"], font_size, cfg["bold"])
                        except Exception:
                            pass

        # 清理空行
        if clear_blank:
            paras = list(doc.paragraphs)
            blank_count =0
            for p in reversed(paras):
                if not p.text.strip():
                    blank_count +=1
                    if blank_count>max_blank:
                        p._element.getparent().remove(p._element)
                else:
                    blank_count=0

        # 保存
        output_path = safe_save_temp_file(b"")
        doc.save(output_path)
        with open(output_path,"rb") as f:
            return f.read(), stats, number_groups

    except Exception as e:
        print(f"处理失败：{str(e)}")
        return None, None, None
    finally:
        safe_delete_file(tmp_path)
        safe_delete_file(output_path)
        clear_memory(doc)